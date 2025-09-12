#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换工具：将 DOC/PDF 转换为 Markdown 格式
支持层级结构识别和 YAML 前导信息生成
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime
import frontmatter
import yaml

# 第三方库导入
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("请安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import pdfplumber
    import PyPDF2
except ImportError:
    print("请安装 PDF 处理库: pip install pdfplumber PyPDF2")
    sys.exit(1)


class DocumentConverter:
    """文档转换器主类"""
    
    def __init__(self, output_dir="converted_markdown"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def extract_metadata_from_filename(self, filename):
        """从文件名提取元数据"""
        name = Path(filename).stem
        metadata = {
            'title': name,
            'filename': filename,
            'converted_at': datetime.now().isoformat(),
            'source_type': 'unknown'
        }
        
        # 提取专业信息
        if '会计' in name:
            metadata['major'] = '会计'
        elif '信管' in name:
            metadata['major'] = '信息管理'
        elif '工商' in name:
            metadata['major'] = '工商管理'
        elif '计科' in name or 'CS' in name:
            metadata['major'] = '计算机科学'
        elif '通信' in name:
            metadata['major'] = '通信工程'
        elif '环境' in name:
            metadata['major'] = '环境工程'
        elif '数媒' in name:
            metadata['major'] = '数字媒体'
        
        # 提取年份
        year_match = re.search(r'(20\d{2})', name)
        if year_match:
            metadata['year'] = year_match.group(1)
        
        # 提取学生姓名（通常在"飞跃手册-"后面）
        name_match = re.search(r'飞跃手册[-_]?(.+)', name)
        if name_match:
            metadata['student_name'] = name_match.group(1).replace('-', ' ').replace('_', ' ')
        
        return metadata
    
    def convert_docx_to_markdown(self, docx_path):
        """将 DOCX 文件转换为 Markdown"""
        try:
            doc = Document(docx_path)
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 检测标题级别（基于字体大小和样式）
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.split()[-1])
                    markdown_content.append(f"{'#' * level} {text}")
                elif paragraph.style.name == 'Title':
                    markdown_content.append(f"# {text}")
                elif paragraph.style.name == 'Subtitle':
                    markdown_content.append(f"## {text}")
                elif len(text) > 50 and not text.endswith(('.', '。', '!', '！', '?', '？')):
                    # 可能是标题，使用二级标题
                    markdown_content.append(f"## {text}")
                else:
                    markdown_content.append(text)
                
                markdown_content.append("")  # 添加空行
            
            # 处理表格
            for table in doc.tables:
                markdown_content.append(self._convert_table_to_markdown(table))
                markdown_content.append("")
            
            return "\n".join(markdown_content)
            
        except Exception as e:
            print(f"转换 DOCX 文件时出错: {e}")
            return f"# 转换错误\n\n无法读取文件: {docx_path}\n错误: {str(e)}"
    
    def convert_pdf_to_markdown(self, pdf_path):
        """将 PDF 文件转换为 Markdown - 基于视觉结构分析"""
        try:
            markdown_content = []
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本和字体信息
                    text_objects = page.chars
                    text_lines = page.extract_text_simple().split('\n')
                    
                    # 分析字体大小分布
                    font_sizes = [char['size'] for char in text_objects if char['size']]
                    if font_sizes:
                        avg_font_size = sum(font_sizes) / len(font_sizes)
                        large_font_threshold = avg_font_size * 1.2
                        small_font_threshold = avg_font_size * 0.8
                    else:
                        avg_font_size = 12
                        large_font_threshold = 14
                        small_font_threshold = 10
                    
                    # 处理每一行文本
                    processed_lines = self._process_text_lines(text_lines, text_objects, 
                                                            large_font_threshold, small_font_threshold)
                    markdown_content.extend(processed_lines)
                    
                    if page_num < len(pdf.pages):
                        markdown_content.append("\n---\n")
            
            return "\n".join(markdown_content)
            
        except Exception as e:
            print(f"转换 PDF 文件时出错: {e}")
            return f"# 转换错误\n\n无法读取文件: {pdf_path}\n错误: {str(e)}"
    
    def _process_text_lines(self, text_lines, text_objects, large_font_threshold, small_font_threshold):
        """处理文本行，基于视觉结构分析"""
        processed_lines = []
        current_paragraph = []
        in_list = False
        
        for line in text_lines:
            line = line.strip()
            if not line:
                # 空行：结束当前段落
                if current_paragraph:
                    processed_lines.extend(self._finalize_paragraph(current_paragraph, in_list))
                    current_paragraph = []
                    in_list = False
                processed_lines.append('')
                continue
            
            # 分析行的特征
            line_features = self._analyze_line_features(line, text_objects)
            
            # 判断是否为标题
            if self._is_visual_title(line, line_features, large_font_threshold):
                # 结束当前段落
                if current_paragraph:
                    processed_lines.extend(self._finalize_paragraph(current_paragraph, in_list))
                    current_paragraph = []
                    in_list = False
                
                # 添加标题
                title_level = self._determine_title_level(line)
                processed_lines.append(f"{'#' * title_level} {line}")
                processed_lines.append('')
            elif self._is_list_item(line):
                # 列表项
                if current_paragraph and not in_list:
                    processed_lines.extend(self._finalize_paragraph(current_paragraph, False))
                    current_paragraph = []
                
                in_list = True
                processed_lines.append(f"- {line}")
            else:
                # 普通文本
                if in_list:
                    # 结束列表
                    processed_lines.append('')
                    in_list = False
                
                current_paragraph.append(line)
        
        # 处理最后一个段落
        if current_paragraph:
            processed_lines.extend(self._finalize_paragraph(current_paragraph, in_list))
        
        return processed_lines
    
    def _analyze_line_features(self, line, text_objects):
        """分析行的视觉特征"""
        # 这里可以添加更复杂的视觉分析
        return {
            'length': len(line),
            'starts_with_number': bool(re.match(r'^\d+[\.\)]', line)),
            'starts_with_bullet': line.startswith(('•', '·', '-', '*')),
            'ends_with_colon': line.endswith(('：', ':')),
            'is_question': line.endswith('?') or line.endswith('？'),
            'is_short': len(line) < 50
        }
    
    def _is_visual_title(self, line, features, large_font_threshold):
        """基于视觉特征判断是否为标题"""
        # 明确的标题模式
        title_patterns = [
            r'^前言$', r'^个人情况$', r'^申请准备$', r'^经验总结$',
            r'^我应该.*\?$', r'^如何.*\?$', r'^.*规划.*\?$',
            r'^GPA[:：]', r'^雅思[:：]', r'^托福[:：]', r'^GRE[:：]',
            r'^去向[:：]', r'^硬背景[:：]', r'^软背景[:：]'
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, line):
                return True
        
        # 短句且以冒号结尾
        if features['is_short'] and features['ends_with_colon']:
            return True
        
        # 数字编号的标题
        if features['starts_with_number'] and features['is_short']:
            return True
        
        return False
    
    def _is_list_item(self, line):
        """判断是否为列表项"""
        return (line.startswith(('•', '·', '-', '*')) or 
                re.match(r'^\d+[\.\)]\s', line) or
                re.match(r'^[一二三四五六七八九十]+[、\.]\s', line))
    
    def _determine_title_level(self, line):
        """确定标题级别"""
        # 主要章节
        if re.match(r'^(前言|个人情况|申请准备|经验总结|最后想说的话)$', line):
            return 1
        
        # 问题标题
        if re.match(r'^我应该.*\?$', line):
            return 1
        
        # 具体主题
        if re.match(r'^(GPA|雅思|托福|GRE|去向|硬背景|软背景)[:：]', line):
            return 2
        
        # 编号标题
        if re.match(r'^\d+\.\s', line):
            return 2
        
        # 默认二级标题
        return 2
    
    def _finalize_paragraph(self, paragraph_lines, was_list):
        """完成段落处理"""
        if not paragraph_lines:
            return []
        
        # 合并段落行
        paragraph_text = ' '.join(paragraph_lines)
        
        # 清理多余的标点符号
        paragraph_text = re.sub(r'\s+([，。；：])', r'\1', paragraph_text)
        paragraph_text = re.sub(r'([，。；：])\s+', r'\1 ', paragraph_text)
        
        return [paragraph_text, '']
    
    def _detect_title_level(self, line, section_stack):
        """智能检测标题层级"""
        # 飞跃手册特有的标题模式
        title_patterns = {
            # 一级标题：主要章节
            1: [
                r'^前言$', r'^个人情况$', r'^申请准备$', r'^申请过程$', 
                r'^经验总结$', r'^建议$', r'^最后想说的话$', r'^结语$',
                r'^我应该.*\?$', r'^如何.*\?$', r'^.*规划.*\?$'
            ],
            # 二级标题：具体主题
            2: [
                r'^GPA[:：]', r'^雅思[:：]', r'^托福[:：]', r'^GRE[:：]', 
                r'^去向[:：]', r'^硬背景[:：]', r'^软背景[:：]', r'^语言[:：]',
                r'^实习[:：]', r'^项目[:：]', r'^科研[:：]', r'^CV[:：]', r'^PS[:：]',
                r'^笔面[:：]', r'^Tips[:：]', r'^时间线[:：]', r'^时间规划[:：]',
                r'^[0-9]+\.\s*[^。]+$',  # 数字编号的标题
                r'^[一二三四五六七八九十]+[、\.]\s*[^。]+$'  # 中文编号
            ],
            # 三级标题：细分内容
            3: [
                r'^[0-9]+\.[0-9]+\s*[^。]+$',  # 二级编号
                r'^[a-zA-Z]\.\s*[^。]+$',  # 字母编号
                r'^[•·]\s*[^。]+$'  # 项目符号
            ]
        }
        
        # 检查各种标题模式
        for level, patterns in title_patterns.items():
            for pattern in patterns:
                if re.match(pattern, line):
                    return level
        
        # 特殊处理：短句且以冒号结尾
        if len(line) < 30 and line.endswith(('：', ':')):
            return 2
        
        # 特殊处理：全大写或首字母大写的短句
        if len(line) < 50 and (line.isupper() or line.istitle()):
            return 2
        
        # 不是标题
        return 0
    
    def _is_title(self, line):
        """简单判断是否为标题（保留兼容性）"""
        return self._detect_title_level(line, []) > 0
    
    def _convert_table_to_markdown(self, table):
        """将表格转换为 Markdown 格式"""
        markdown_table = []
        
        for i, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            
            markdown_table.append("| " + " | ".join(row_data) + " |")
            
            # 添加表头分隔线
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(row_data)) + " |"
                markdown_table.append(separator)
        
        return "\n".join(markdown_table)
    
    def add_yaml_frontmatter(self, content, metadata):
        """添加 YAML 前导信息"""
        yaml_content = yaml.dump(metadata, default_flow_style=False, allow_unicode=True)
        return f"---\n{yaml_content}---\n\n{content}"
    
    def _post_process_content(self, content):
        """后处理内容，优化结构"""
        lines = content.split('\n')
        processed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                processed_lines.append('')
                i += 1
                continue
            
            # 处理连续的空行
            if processed_lines and processed_lines[-1] == '' and line == '':
                i += 1
                continue
            
            # 处理标题后的空行
            if line.startswith('#'):
                processed_lines.append(line)
                # 确保标题后有适当的空行
                if i + 1 < len(lines) and lines[i + 1].strip():
                    processed_lines.append('')
            else:
                processed_lines.append(line)
            
            i += 1
        
        # 清理多余的空行
        result_lines = []
        prev_empty = False
        
        for line in processed_lines:
            if line == '':
                if not prev_empty:
                    result_lines.append(line)
                prev_empty = True
            else:
                result_lines.append(line)
                prev_empty = False
        
        return '\n'.join(result_lines)
    
    def convert_file(self, file_path):
        """转换单个文件"""
        file_path = Path(file_path)
        if not file_path.exists():
            print(f"文件不存在: {file_path}")
            return None
        
        # 提取元数据
        metadata = self.extract_metadata_from_filename(file_path.name)
        
        # 根据文件类型选择转换方法
        if file_path.suffix.lower() in ['.docx', '.doc']:
            metadata['source_type'] = 'docx'
            content = self.convert_docx_to_markdown(file_path)
        elif file_path.suffix.lower() == '.pdf':
            metadata['source_type'] = 'pdf'
            content = self.convert_pdf_to_markdown(file_path)
        else:
            print(f"不支持的文件类型: {file_path.suffix}")
            return None
        
        # 后处理内容
        content = self._post_process_content(content)
        
        # 添加 YAML 前导信息
        full_content = self.add_yaml_frontmatter(content, metadata)
        
        # 保存转换后的文件
        output_filename = file_path.stem + '.md'
        output_path = self.output_dir / output_filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        print(f"✅ 转换完成: {file_path.name} -> {output_filename}")
        return output_path
    
    def batch_convert(self, directory_path, recursive=True):
        """批量转换目录中的文件"""
        directory_path = Path(directory_path)
        if not directory_path.exists():
            print(f"目录不存在: {directory_path}")
            return
        
        supported_extensions = ['.docx', '.doc', '.pdf']
        converted_files = []
        failed_files = []
        
        # 查找所有支持的文件
        if recursive:
            files = []
            for ext in supported_extensions:
                files.extend(directory_path.rglob(f"*{ext}"))
        else:
            files = []
            for ext in supported_extensions:
                files.extend(directory_path.glob(f"*{ext}"))
        
        print(f"找到 {len(files)} 个文件需要转换...")
        
        for file_path in files:
            try:
                result = self.convert_file(file_path)
                if result:
                    converted_files.append(result)
                else:
                    failed_files.append(file_path)
            except Exception as e:
                print(f"❌ 转换失败: {file_path.name} - {e}")
                failed_files.append(file_path)
        
        print(f"\n📊 转换统计:")
        print(f"✅ 成功: {len(converted_files)} 个文件")
        print(f"❌ 失败: {len(failed_files)} 个文件")
        
        if failed_files:
            print(f"\n失败文件列表:")
            for file_path in failed_files:
                print(f"  - {file_path}")
        
        return converted_files, failed_files


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='文档转换工具：将 DOC/PDF 转换为 Markdown')
    parser.add_argument('input', help='输入文件或目录路径')
    parser.add_argument('-o', '--output', default='converted_markdown', 
                       help='输出目录 (默认: converted_markdown)')
    parser.add_argument('-r', '--recursive', action='store_true', 
                       help='递归处理子目录')
    
    args = parser.parse_args()
    
    converter = DocumentConverter(args.output)
    
    input_path = Path(args.input)
    
    if input_path.is_file():
        # 转换单个文件
        converter.convert_file(input_path)
    elif input_path.is_dir():
        # 批量转换目录
        converter.batch_convert(input_path, recursive=args.recursive)
    else:
        print(f"路径不存在: {input_path}")


if __name__ == "__main__":
    main()